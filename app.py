from flask import Flask, render_template, request, redirect, url_for, flash, session, jsonify
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime, timedelta
import os
import socket
from dotenv import load_dotenv
import telepot
import schedule
import time
import threading
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from datetime import datetime
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import io
import base64
import logging
from logging.handlers import RotatingFileHandler

# Load environment variables from static folder
load_dotenv('static/.env')

# Telegram credentials
TELEGRAM_TOKEN = "8400465991:AAGr7sjaRqL5T6Xq_UoAW0_5qTn5jAPvrvY"
TELEGRAM_ID = "-1002961398669"

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-here'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///ases_stok.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Log sistemi konfigürasyonu
def setup_logging():
    """Log sistemi kurulumu"""
    # Log klasörünü oluştur
    import os
    if not os.path.exists('logs'):
        os.makedirs('logs')
    
    # Ana logger
    app.logger.setLevel(logging.INFO)
    
    # Dosya handler - tüm loglar
    file_handler = RotatingFileHandler('logs/ases_stok.log', maxBytes=10240000, backupCount=10)
    file_handler.setFormatter(logging.Formatter(
        '%(asctime)s %(levelname)s: %(message)s [in %(pathname)s:%(lineno)d]'
    ))
    file_handler.setLevel(logging.INFO)
    app.logger.addHandler(file_handler)
    
    # Hareket logları için ayrı handler
    movement_handler = RotatingFileHandler('logs/movements.log', maxBytes=10240000, backupCount=10)
    movement_handler.setFormatter(logging.Formatter(
        '%(asctime)s - %(message)s'
    ))
    movement_handler.setLevel(logging.INFO)
    
    # Hareket logger'ı
    movement_logger = logging.getLogger('movements')
    movement_logger.addHandler(movement_handler)
    movement_logger.setLevel(logging.INFO)
    
    # Veritabanı değişiklikleri için ayrı handler
    db_handler = RotatingFileHandler('logs/database.log', maxBytes=10240000, backupCount=10)
    db_handler.setFormatter(logging.Formatter(
        '%(asctime)s - %(message)s'
    ))
    db_handler.setLevel(logging.INFO)
    
    # Veritabanı logger'ı
    db_logger = logging.getLogger('database')
    db_logger.addHandler(db_handler)
    db_logger.setLevel(logging.INFO)
    
    app.logger.info('Ases Stok Takip sistemi başlatıldı')

# Log sistemi kurulumu
setup_logging()

db = SQLAlchemy(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

# Veritabanı Modelleri
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password_hash = db.Column(db.String(120), nullable=False)
    is_admin = db.Column(db.Boolean, default=False)
    
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class Product(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    marka = db.Column(db.String(100), nullable=False)
    model = db.Column(db.String(100), nullable=False)
    barkod = db.Column(db.String(100), nullable=False)
    seri_no = db.Column(db.String(100), unique=True, nullable=False)
    adet = db.Column(db.Integer, default=0)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class Movement(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    product_id = db.Column(db.Integer, db.ForeignKey('product.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    movement_type = db.Column(db.String(20), nullable=False)  # 'giris' or 'cikis'
    adet = db.Column(db.Integer, nullable=False)
    aciklama = db.Column(db.Text)
    tarih = db.Column(db.DateTime, default=datetime.utcnow)
    
    product = db.relationship('Product', backref=db.backref('movements', lazy=True))
    user = db.relationship('User', backref=db.backref('movements', lazy=True))

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Telegram notification function
def send_telegram_notification(movement_type, product_info, user_info, adet, aciklama=""):
    """Send movement notification to Telegram"""
    try:
        telegram_id = int(TELEGRAM_ID)
        token = TELEGRAM_TOKEN
        
        if not telegram_id or not token or token == "YENİ_BOT_TOKEN_BURAYA":
            print("Telegram credentials not found or not updated")
            return False
            
        bot = telepot.Bot(token)
        
        send_time = datetime.now()
        movement_emoji = "📥" if movement_type == "giris" else "📤"
        movement_text = "GİRİŞ" if movement_type == "giris" else "ÇIKIŞ"
        
        message_text = f"{movement_emoji} STOK HAREKETİ TESPİT EDİLDİ!\n"
        message_text += f"🗓 Tarih: {send_time.strftime('%d.%m.%Y')}\n"
        message_text += f"⏰ Saat: {send_time.strftime('%H:%M:%S')}\n"
        message_text += f"📦 Hareket Türü: {movement_text}\n"
        message_text += f"🏷️ Marka: {product_info['marka']}\n"
        message_text += f"📱 Model: {product_info['model']}\n"
        message_text += f"🔢 Barkod: {product_info['barkod']}\n"
        message_text += f"📊 Adet: {adet}\n"
        message_text += f"👤 Kullanıcı: {user_info['username']}\n"
        message_text += f"💻 Cihaz: {socket.gethostname()}\n"
        
        if aciklama:
            message_text += f"📝 Açıklama: {aciklama}\n"
        
        bot.sendMessage(telegram_id, message_text)
        print(f"Telegram notification sent for {movement_type} movement")
        return True
        
    except Exception as e:
        print(f"Error sending Telegram notification: {e}")
        return False

# Excel rapor oluşturma fonksiyonu
def create_stock_excel_report():
    """Güncel stok durumunu Excel dosyası olarak oluştur"""
    try:
        with app.app_context():
            # Tüm ürünleri al
            products = Product.query.all()
            
            # Excel workbook oluştur
            wb = Workbook()
            ws = wb.active
            ws.title = "Stok Durumu"
            
            # Başlık satırı
            headers = ['ID', 'Marka', 'Model', 'Barkod', 'Seri No', 'Adet', 'Oluşturma Tarihi']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                cell.alignment = Alignment(horizontal="center")
            
            # Veri satırları
            for row, product in enumerate(products, 2):
                ws.cell(row=row, column=1, value=product.id)
                ws.cell(row=row, column=2, value=product.marka)
                ws.cell(row=row, column=3, value=product.model)
                ws.cell(row=row, column=4, value=product.barkod)
                ws.cell(row=row, column=5, value=product.seri_no)
                ws.cell(row=row, column=6, value=product.adet)
                ws.cell(row=row, column=7, value=product.created_at.strftime('%d.%m.%Y %H:%M'))
            
            # Sütun genişliklerini ayarla
            column_widths = [8, 15, 20, 15, 15, 8, 18]
            for col, width in enumerate(column_widths, 1):
                ws.column_dimensions[ws.cell(row=1, column=col).column_letter].width = width
            
            # Dosya adı oluştur
            today = datetime.now().strftime('%Y%m%d')
            filename = f"stok_durumu_{today}.xlsx"
            filepath = os.path.join('static', filename)
            
            # Excel dosyasını kaydet
            wb.save(filepath)
            print(f"Excel raporu oluşturuldu: {filename}")
            return filepath, filename
            
    except Exception as e:
        print(f"Excel raporu oluşturulurken hata: {e}")
        return None, None

# Telegram'a Excel dosyası gönderme fonksiyonu
def send_excel_via_telegram(filepath, filename):
    """Excel dosyasını Telegram'a gönder"""
    try:
        telegram_id = int(TELEGRAM_ID)
        token = TELEGRAM_TOKEN
        
        if not telegram_id or not token or token == "YENİ_BOT_TOKEN_BURAYA":
            print("Telegram credentials not found or not updated")
            return False
            
        bot = telepot.Bot(token)
        
        # Mesaj metni
        today = datetime.now().strftime('%d.%m.%Y')
        message_text = f"📊 GÜNLÜK STOK RAPORU\n"
        message_text += f"📅 Tarih: {today}\n"
        message_text += f"⏰ Saat: {datetime.now().strftime('%H:%M')}\n"
        message_text += f"💻 Cihaz: {socket.gethostname()}\n"
        message_text += f"📁 Dosya: {filename}"
        
        # Excel dosyasını gönder
        with open(filepath, 'rb') as file:
            bot.sendDocument(telegram_id, file, caption=message_text)
        
        print(f"Excel raporu Telegram'a gönderildi: {filename}")
        return True
        
    except Exception as e:
        print(f"Excel dosyası Telegram'a gönderilirken hata: {e}")
        return False

# Veritabanı yedekleme fonksiyonu
def create_database_backup():
    """Veritabanını yedekle"""
    try:
        # Mevcut veritabanı dosyası
        source_db = os.path.join('instance', 'ases_stok.db')
        
        # Yedek dosya adı
        today = datetime.now().strftime('%Y%m%d')
        backup_filename = f"ases_stok_backup_{today}.db"
        backup_filepath = os.path.join('static', backup_filename)
        
        # Dosyayı kopyala
        import shutil
        shutil.copy2(source_db, backup_filepath)
        
        print(f"Veritabanı yedeklendi: {backup_filename}")
        return backup_filepath, backup_filename
        
    except Exception as e:
        print(f"Veritabanı yedeklenirken hata: {e}")
        return None, None

# Telegram'a veritabanı gönderme fonksiyonu
def send_database_via_telegram(filepath, filename):
    """Veritabanı dosyasını Telegram'a gönder"""
    try:
        telegram_id = int(TELEGRAM_ID)
        token = TELEGRAM_TOKEN
        
        if not telegram_id or not token or token == "YENİ_BOT_TOKEN_BURAYA":
            print("Telegram credentials not found or not updated")
            return False
            
        bot = telepot.Bot(token)
        
        # Mesaj metni
        today = datetime.now().strftime('%d.%m.%Y')
        message_text = f"💾 VERİTABANI YEDEĞİ\n"
        message_text += f"📅 Tarih: {today}\n"
        message_text += f"⏰ Saat: {datetime.now().strftime('%H:%M')}\n"
        message_text += f"💻 Cihaz: {socket.gethostname()}\n"
        message_text += f"📁 Dosya: {filename}\n"
        message_text += f"📊 Boyut: {os.path.getsize(filepath)} bytes"
        
        # Veritabanı dosyasını gönder
        with open(filepath, 'rb') as file:
            bot.sendDocument(telegram_id, file, caption=message_text)
        
        print(f"Veritabanı yedeği Telegram'a gönderildi: {filename}")
        return True
        
    except Exception as e:
        print(f"Veritabanı dosyası Telegram'a gönderilirken hata: {e}")
        return False

# Günlük stok raporu ve veritabanı yedekleme fonksiyonu
def send_daily_stock_report():
    """Günlük stok raporunu ve veritabanı yedeğini oluştur ve gönder"""
    print("Günlük rapor ve yedekleme işlemi başlatılıyor...")
    
    with app.app_context():
        # Excel raporu oluştur
        excel_filepath, excel_filename = create_stock_excel_report()
        
        # Veritabanı yedeği oluştur
        db_filepath, db_filename = create_database_backup()
        
        # Excel raporunu gönder
        if excel_filepath and excel_filename:
            success_excel = send_excel_via_telegram(excel_filepath, excel_filename)
            if success_excel:
                print(f"Excel dosyası korundu: {excel_filepath}")
            else:
                print("Excel raporu Telegram'a gönderim başarısız")
        else:
            print("Excel raporu oluşturulamadı")
        
        # Veritabanı yedeğini gönder
        if db_filepath and db_filename:
            success_db = send_database_via_telegram(db_filepath, db_filename)
            if success_db:
                print(f"Veritabanı yedeği korundu: {db_filepath}")
            else:
                print("Veritabanı yedeği Telegram'a gönderim başarısız")
        else:
            print("Veritabanı yedeği oluşturulamadı")
    
    print("Günlük yedekleme işlemi tamamlandı!")

# Zamanlayıcı fonksiyonu
def run_scheduler():
    """Zamanlayıcıyı çalıştır"""
    while True:
        schedule.run_pending()
        time.sleep(60)  # Her dakika kontrol et

# Ana Sayfa
@app.route('/')
@login_required
def index():
    # İstatistikler (sadece stokta olan ürünler)
    total_products = Product.query.filter(Product.adet > 0).count()
    total_movements = Movement.query.count()
    total_quantity = db.session.query(db.func.sum(Product.adet)).filter(Product.adet > 0).scalar() or 0
    
    # Son hareketler
    recent_movements = Movement.query.order_by(Movement.tarih.desc()).limit(5).all()
    
    return render_template('index.html', 
                         total_products=total_products,
                         total_movements=total_movements,
                         total_quantity=total_quantity,
                         recent_movements=recent_movements)

# Giriş Sayfası
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = User.query.filter_by(username=username).first()
        
        if user and check_password_hash(user.password_hash, password):
            login_user(user)
            return redirect(url_for('index'))
        else:
            flash('Geçersiz kullanıcı adı veya şifre!', 'error')
    
    return render_template('login.html')

# Çıkış
@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

# Ürün Yönetimi
@app.route('/products')
@login_required
def products():
    search = request.args.get('search', '')
    brand = request.args.get('brand', '')
    model = request.args.get('model', '')
    
    query = Product.query.filter(Product.adet > 0)  # Sadece stokta olan ürünler
    if search:
        query = query.filter(
            db.or_(
                Product.marka.contains(search),
                Product.model.contains(search),
                Product.barkod.contains(search),
                Product.seri_no.contains(search)
            )
        )
    if brand:
        query = query.filter(Product.marka == brand)
    if model:
        query = query.filter(Product.model == model)
    
    products = query.all()
    
    # Filtreleme için marka ve model listelerini al (sadece stokta olan ürünlerden)
    brands = db.session.query(Product.marka).filter(Product.adet > 0).distinct().all()
    models = db.session.query(Product.model).filter(Product.adet > 0).distinct().all()
    
    return render_template('products.html', 
                         products=products, 
                         brands=[b[0] for b in brands],
                         models=[m[0] for m in models])

# Ürün Ekleme
@app.route('/add_product', methods=['GET', 'POST'])
@login_required
def add_product():
    if request.method == 'POST':
        marka = request.form['marka']
        model = request.form['model']
        barkod = request.form['barkod']
        seri_no = request.form['seri_no']
        adet = int(request.form['adet'])
        
        # Sadece seri no benzersizlik kontrolü
        if Product.query.filter_by(seri_no=seri_no).first():
            flash('Bu seri numarası zaten mevcut!', 'error')
            return render_template('add_product.html')
        
        product = Product(
            marka=marka,
            model=model,
            barkod=barkod,
            seri_no=seri_no,
            adet=adet
        )
        
        db.session.add(product)
        db.session.commit()
        
        # Giriş hareketi kaydet
        movement = Movement(
            product_id=product.id,
            user_id=current_user.id,
            movement_type='giris',
            adet=adet,
            aciklama=f'İlk stok girişi - {adet} adet'
        )
        db.session.add(movement)
        db.session.commit()
        
        # Hareket logunu kaydet
        movement_logger = logging.getLogger('movements')
        movement_logger.info(f'GİRİŞ HAREKETİ - {product.marka} {product.model} | {adet} adet | {current_user.username}')
        
        # Veritabanı değişikliğini logla
        db_logger = logging.getLogger('database')
        db_logger.info(f'GİRİŞ HAREKETİ EKLENDİ - Ürün: {product.marka} {product.model} (ID: {product.id})')
        db_logger.info(f'  Hareket: GİRİŞ, {adet} adet, Açıklama: İlk stok girişi - {adet} adet')
        db_logger.info(f'  Stok: 0 → {product.adet}')
        db_logger.info(f'  Kullanıcı: {current_user.username} (ID: {current_user.id})')
        
        # Telegram bildirimi gönder
        product_info = {
            'marka': product.marka,
            'model': product.model,
            'barkod': product.barkod
        }
        user_info = {
            'username': current_user.username,
            'email': current_user.email
        }
        send_telegram_notification('giris', product_info, user_info, adet, f'İlk stok girişi - {adet} adet')
        
        flash('Ürün başarıyla eklendi!', 'success')
        app.logger.info(f'Ürün eklendi: {product.marka} {product.model} - {adet} adet by {current_user.username}')
        return redirect(url_for('products'))
    
    return render_template('add_product.html')

# Hareket Yönetimi
@app.route('/movements')
@login_required
def movements():
    # Tarih filtreleme parametreleri
    date_filter = request.args.get('date_filter', 'all')  # all, today, week, month, year, custom
    custom_date = request.args.get('custom_date', '')
    
    query = Movement.query
    
    # Tarih filtreleme
    if date_filter == 'today':
        today = datetime.now().date()
        query = query.filter(db.func.date(Movement.tarih) == today)
    elif date_filter == 'week':
        week_ago = datetime.now().date() - timedelta(days=7)
        query = query.filter(db.func.date(Movement.tarih) >= week_ago)
    elif date_filter == 'month':
        month_ago = datetime.now().date() - timedelta(days=30)
        query = query.filter(db.func.date(Movement.tarih) >= month_ago)
    elif date_filter == 'year':
        year_ago = datetime.now().date() - timedelta(days=365)
        query = query.filter(db.func.date(Movement.tarih) >= year_ago)
    elif date_filter == 'custom' and custom_date:
        try:
            custom_date_obj = datetime.strptime(custom_date, '%Y-%m-%d').date()
            query = query.filter(db.func.date(Movement.tarih) == custom_date_obj)
        except ValueError:
            pass  # Geçersiz tarih formatı, tüm hareketleri göster
    
    movements = query.order_by(Movement.tarih.desc()).all()
    
    # Tarih istatistikleri
    total_movements = Movement.query.count()
    today_movements = Movement.query.filter(db.func.date(Movement.tarih) == datetime.now().date()).count()
    week_movements = Movement.query.filter(db.func.date(Movement.tarih) >= datetime.now().date() - timedelta(days=7)).count()
    month_movements = Movement.query.filter(db.func.date(Movement.tarih) >= datetime.now().date() - timedelta(days=30)).count()
    
    return render_template('movements.html', 
                         movements=movements,
                         date_filter=date_filter,
                         custom_date=custom_date,
                         total_movements=total_movements,
                         today_movements=today_movements,
                         week_movements=week_movements,
                         month_movements=month_movements)

# Hareket Düzenleme
@app.route('/edit_movement/<int:movement_id>', methods=['GET', 'POST'])
@login_required
def edit_movement(movement_id):
    movement = Movement.query.get_or_404(movement_id)
    
    if request.method == 'POST':
        # Eski değerleri kaydet (stok güncellemesi için)
        old_adet = movement.adet
        old_type = movement.movement_type
        old_aciklama = movement.aciklama
        old_tarih = movement.tarih
        
        # Form verilerini al
        movement.movement_type = request.form['movement_type']
        movement.adet = int(request.form['adet'])
        movement.aciklama = request.form.get('aciklama', '')
        movement.tarih = datetime.strptime(request.form['tarih'], '%Y-%m-%dT%H:%M')
        
        # Stok güncellemesi
        product = movement.product
        
        # Eski hareketi geri al
        if old_type == 'giris':
            product.adet -= old_adet
        else:  # cikis
            product.adet += old_adet
        
        # Yeni hareketi uygula
        if movement.movement_type == 'giris':
            product.adet += movement.adet
        else:  # cikis
            product.adet -= movement.adet
        
        # Negatif stok kontrolü
        if product.adet < 0:
            flash('Stok miktarı negatif olamaz!', 'error')
            app.logger.warning(f'Hareket düzenleme başarısız - Negatif stok: Movement ID {movement_id}, Kullanıcı: {current_user.username}')
            return redirect(url_for('movements'))
        
        # Veritabanı değişikliğini logla
        db_logger = logging.getLogger('database')
        db_logger.info(f'HAREKET DÜZENLENDİ - ID: {movement_id}, Ürün: {product.marka} {product.model} (ID: {product.id})')
        db_logger.info(f'  Eski: {old_type}, {old_adet} adet, {old_tarih.strftime("%d.%m.%Y %H:%M")}, Açıklama: {old_aciklama or "Yok"}')
        db_logger.info(f'  Yeni: {movement.movement_type}, {movement.adet} adet, {movement.tarih.strftime("%d.%m.%Y %H:%M")}, Açıklama: {movement.aciklama or "Yok"}')
        db_logger.info(f'  Stok değişimi: {product.adet - (old_adet if old_type == "giris" else -old_adet)} → {product.adet}')
        db_logger.info(f'  Kullanıcı: {current_user.username} (ID: {current_user.id})')
        
        # Hareket logunu güncelle
        movement_logger = logging.getLogger('movements')
        movement_logger.info(f'HAREKET DÜZENLENDİ - {product.marka} {product.model} | {movement.movement_type.upper()} | {movement.adet} adet | {current_user.username}')
        
        db.session.commit()
        flash('Hareket başarıyla güncellendi!', 'success')
        app.logger.info(f'Hareket başarıyla düzenlendi: ID {movement_id} by {current_user.username}')
        return redirect(url_for('movements'))
    
    return render_template('edit_movement.html', movement=movement)

# Hareket Silme
@app.route('/delete_movement/<int:movement_id>')
@login_required
def delete_movement(movement_id):
    movement = Movement.query.get_or_404(movement_id)
    product = movement.product
    
    # Hareket bilgilerini logla
    movement_logger = logging.getLogger('movements')
    movement_logger.info(f'HAREKET SİLİNDİ - {product.marka} {product.model} | {movement.movement_type.upper()} | {movement.adet} adet | {current_user.username}')
    
    # Veritabanı değişikliğini logla
    db_logger = logging.getLogger('database')
    db_logger.info(f'HAREKET SİLİNDİ - ID: {movement_id}, Ürün: {product.marka} {product.model} (ID: {product.id})')
    db_logger.info(f'  Hareket: {movement.movement_type}, {movement.adet} adet, {movement.tarih.strftime("%d.%m.%Y %H:%M")}, Açıklama: {movement.aciklama or "Yok"}')
    db_logger.info(f'  Kullanıcı: {current_user.username} (ID: {current_user.id})')
    
    # Stok güncellemesi - hareketi geri al
    if movement.movement_type == 'giris':
        product.adet -= movement.adet
    else:  # cikis
        product.adet += movement.adet
    
    # Negatif stok kontrolü
    if product.adet < 0:
        flash('Bu hareketi silmek stok miktarını negatif yapar!', 'error')
        app.logger.warning(f'Hareket silme başarısız - Negatif stok: Movement ID {movement_id}, Kullanıcı: {current_user.username}')
        return redirect(url_for('movements'))
    
    db_logger.info(f'  Stok değişimi: {product.adet + (movement.adet if movement.movement_type == "giris" else -movement.adet)} → {product.adet}')
    
    db.session.delete(movement)
    db.session.commit()
    flash('Hareket başarıyla silindi!', 'success')
    app.logger.info(f'Hareket başarıyla silindi: ID {movement_id} by {current_user.username}')
    return redirect(url_for('movements'))

# Çıkış Hareketi
@app.route('/exit_product', methods=['GET', 'POST'])
@login_required
def exit_product():
    if request.method == 'POST':
        product_id = request.form['product_id']
        adet = int(request.form['adet'])
        aciklama = request.form['aciklama']
        
        product = Product.query.get(product_id)
        if not product:
            flash('Ürün bulunamadı!', 'error')
            return redirect(url_for('exit_product'))
        
        if product.adet < adet:
            flash('Yetersiz stok!', 'error')
            return redirect(url_for('exit_product'))
        
        # Stoktan düş
        product.adet -= adet
        
        # Çıkış hareketi kaydet
        movement = Movement(
            product_id=product.id,
            user_id=current_user.id,
            movement_type='cikis',
            adet=adet,
            aciklama=aciklama
        )
        
        db.session.add(movement)
        db.session.commit()
        
        # Hareket logunu kaydet
        movement_logger = logging.getLogger('movements')
        movement_logger.info(f'ÇIKIŞ HAREKETİ - {product.marka} {product.model} | {adet} adet | {current_user.username}')
        
        # Veritabanı değişikliğini logla
        db_logger = logging.getLogger('database')
        db_logger.info(f'ÇIKIŞ HAREKETİ EKLENDİ - Ürün: {product.marka} {product.model} (ID: {product.id})')
        db_logger.info(f'  Hareket: ÇIKIŞ, {adet} adet, Açıklama: {aciklama or "Yok"}')
        db_logger.info(f'  Stok değişimi: {product.adet + adet} → {product.adet}')
        db_logger.info(f'  Kullanıcı: {current_user.username} (ID: {current_user.id})')
        
        # Telegram bildirimi gönder
        product_info = {
            'marka': product.marka,
            'model': product.model,
            'barkod': product.barkod
        }
        user_info = {
            'username': current_user.username,
            'email': current_user.email
        }
        send_telegram_notification('cikis', product_info, user_info, adet, aciklama)
        
        flash('Çıkış hareketi başarıyla kaydedildi!', 'success')
        app.logger.info(f'Çıkış hareketi eklendi: {product.marka} {product.model} - {adet} adet by {current_user.username}')
        return redirect(url_for('movements'))
    
    products = Product.query.filter(Product.adet > 0).all()
    return render_template('exit_product.html', products=products)

# Kullanıcı Yönetimi
@app.route('/users')
@login_required
def users():
    if not current_user.is_admin:
        flash('Bu sayfaya erişim yetkiniz yok!', 'error')
        return redirect(url_for('index'))
    
    users = User.query.all()
    return render_template('users.html', users=users)

# Kullanıcı Ekleme
@app.route('/add_user', methods=['GET', 'POST'])
@login_required
def add_user():
    if not current_user.is_admin:
        flash('Bu sayfaya erişim yetkiniz yok!', 'error')
        return redirect(url_for('index'))
    
    if request.method == 'POST':
        username = request.form['username']
        email = request.form['email']
        password = request.form['password']
        is_admin = 'is_admin' in request.form
        
        if User.query.filter_by(email=email).first():
            flash('Bu email zaten kullanılıyor!', 'error')
            return render_template('add_user.html')
        
        user = User(
            username=username,
            email=email,
            password_hash=generate_password_hash(password),
            is_admin=is_admin
        )
        
        db.session.add(user)
        db.session.commit()
        
        flash('Kullanıcı başarıyla eklendi!', 'success')
        return redirect(url_for('users'))
    
    return render_template('add_user.html')

# Kullanıcı Silme
@app.route('/delete_user/<int:user_id>')
@login_required
def delete_user(user_id):
    if not current_user.is_admin:
        flash('Bu işlem için yetkiniz yok!', 'error')
        return redirect(url_for('users'))
    
    user = User.query.get(user_id)
    if user and user.id != current_user.id:
        db.session.delete(user)
        db.session.commit()
        flash('Kullanıcı silindi!', 'success')
    else:
        flash('Kendinizi silemezsiniz!', 'error')
    
    return redirect(url_for('users'))

# Analiz Sayfası
@app.route('/analytics')
@login_required
def analytics():
    # Tarih bazlı analiz
    today = datetime.now().date()
    week_ago = today - timedelta(days=7)
    month_ago = today - timedelta(days=30)
    
    # Günlük hareketler
    daily_movements_raw = db.session.query(
        db.func.date(Movement.tarih).label('date'),
        Movement.movement_type,
        db.func.sum(Movement.adet).label('total')
    ).filter(
        Movement.tarih >= week_ago
    ).group_by(
        db.func.date(Movement.tarih),
        Movement.movement_type
    ).all()
    
    # Convert Row objects to dictionaries for JSON serialization
    daily_movements = []
    for movement in daily_movements_raw:
        daily_movements.append({
            'date': str(movement.date),
            'movement_type': movement.movement_type,
            'total': movement.total
        })
    
    # Toplam istatistikler (sadece stokta olan ürünler)
    total_products = Product.query.filter(Product.adet > 0).count()
    total_quantity = db.session.query(db.func.sum(Product.adet)).filter(Product.adet > 0).scalar() or 0
    avg_stock = total_quantity / total_products if total_products > 0 else 0
    
    # Stok durumu analizi (sadece stokta olan ürünler)
    high_stock = Product.query.filter(Product.adet > 50).count()
    medium_stock = Product.query.filter(Product.adet.between(10, 50)).count()
    low_stock = Product.query.filter(Product.adet.between(1, 9)).count()
    no_stock = Product.query.filter(Product.adet == 0).count()
    
    # Marka bazlı dağılım (sadece stokta olan ürünler)
    brand_distribution = db.session.query(
        Product.marka,
        db.func.count(Product.id).label('product_count'),
        db.func.sum(Product.adet).label('total_quantity')
    ).filter(Product.adet > 0).group_by(Product.marka).order_by(db.func.sum(Product.adet).desc()).limit(10).all()
    
    brand_data = []
    for brand in brand_distribution:
        brand_data.append({
            'marka': brand.marka,
            'product_count': brand.product_count,
            'total_quantity': brand.total_quantity
        })
    
    # En çok hareket gören ürünler (sadece stokta olan ürünler)
    top_products = db.session.query(
        Product.marka,
        Product.model,
        db.func.count(Movement.id).label('movement_count'),
        db.func.sum(Movement.adet).label('total_movement')
    ).join(Movement).filter(Product.adet > 0).group_by(
        Product.id, Product.marka, Product.model
    ).order_by(db.func.count(Movement.id).desc()).limit(10).all()
    
    top_products_data = []
    for product in top_products:
        top_products_data.append({
            'marka': product.marka,
            'model': product.model,
            'movement_count': product.movement_count,
            'total_movement': product.total_movement
        })
    
    # Kullanıcı aktivite analizi
    user_activity = db.session.query(
        User.username,
        db.func.count(Movement.id).label('movement_count'),
        db.func.sum(Movement.adet).label('total_movement')
    ).join(Movement).group_by(User.id, User.username).order_by(
        db.func.count(Movement.id).desc()
    ).all()
    
    user_activity_data = []
    for user in user_activity:
        user_activity_data.append({
            'username': user.username,
            'movement_count': user.movement_count,
            'total_movement': user.total_movement
        })
    
    # Aylık trend analizi
    monthly_trend = db.session.query(
        db.func.strftime('%Y-%m', Movement.tarih).label('month'),
        Movement.movement_type,
        db.func.sum(Movement.adet).label('total')
    ).filter(
        Movement.tarih >= month_ago
    ).group_by(
        db.func.strftime('%Y-%m', Movement.tarih),
        Movement.movement_type
    ).all()
    
    monthly_data = []
    for trend in monthly_trend:
        monthly_data.append({
            'month': trend.month,
            'movement_type': trend.movement_type,
            'total': trend.total
        })
    
    return render_template('analytics.html',
                         daily_movements=daily_movements,
                         total_products=total_products,
                         total_quantity=total_quantity,
                         avg_stock=round(avg_stock, 2),
                         high_stock=high_stock,
                         medium_stock=medium_stock,
                         low_stock=low_stock,
                         no_stock=no_stock,
                         brand_data=brand_data,
                         top_products_data=top_products_data,
                         user_activity_data=user_activity_data,
                         monthly_data=monthly_data)

# Rapor oluşturma fonksiyonları
def create_word_report():
    """Word formatında rapor oluştur"""
    doc = Document()
    
    # Başlık
    title = doc.add_heading('ASES STOK TAKİP SİSTEMİ - ANALİZ RAPORU', 0)
    title.alignment = 1  # Ortala
    
    # Tarih
    doc.add_paragraph(f'Rapor Tarihi: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    doc.add_paragraph('')
    
    # Genel Bilgiler
    doc.add_heading('1. GENEL BİLGİLER', level=1)
    
    # İstatistikler
    total_products = Product.query.count()
    total_quantity = db.session.query(db.func.sum(Product.adet)).scalar() or 0
    avg_stock = total_quantity / total_products if total_products > 0 else 0
    
    # Stok durumu
    high_stock = Product.query.filter(Product.adet > 50).count()
    medium_stock = Product.query.filter(Product.adet.between(10, 50)).count()
    low_stock = Product.query.filter(Product.adet.between(1, 9)).count()
    no_stock = Product.query.filter(Product.adet == 0).count()
    
    stats_data = [
        ['Toplam Ürün Sayısı', str(total_products)],
        ['Toplam Stok Miktarı', str(total_quantity)],
        ['Ortalama Stok', f"{avg_stock:.2f}"],
        ['Yüksek Stok (>50)', str(high_stock)],
        ['Orta Stok (10-50)', str(medium_stock)],
        ['Düşük Stok (1-9)', str(low_stock)],
        ['Stok Yok (0)', str(no_stock)]
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Kategori'
    hdr_cells[1].text = 'Değer'
    
    for stat in stats_data:
        row_cells = table.add_row().cells
        row_cells[0].text = stat[0]
        row_cells[1].text = stat[1]
    
    doc.add_paragraph('')
    
    # Son 7 günlük hareketler
    doc.add_heading('2. SON 7 GÜNLÜK HAREKETLER', level=1)
    
    today = datetime.now().date()
    week_ago = today - timedelta(days=7)
    
    movements = db.session.query(
        db.func.date(Movement.tarih).label('date'),
        Movement.movement_type,
        db.func.sum(Movement.adet).label('total')
    ).filter(
        Movement.tarih >= week_ago
    ).group_by(
        db.func.date(Movement.tarih),
        Movement.movement_type
    ).order_by(db.func.date(Movement.tarih)).all()
    
    if movements:
        movement_table = doc.add_table(rows=1, cols=3)
        movement_table.style = 'Table Grid'
        hdr_cells = movement_table.rows[0].cells
        hdr_cells[0].text = 'Tarih'
        hdr_cells[1].text = 'Hareket Türü'
        hdr_cells[2].text = 'Adet'
        
        for movement in movements:
            row_cells = movement_table.add_row().cells
            row_cells[0].text = str(movement.date)
            row_cells[1].text = 'Giriş' if movement.movement_type == 'giris' else 'Çıkış'
            row_cells[2].text = str(movement.total)
    else:
        doc.add_paragraph('Son 7 günde hareket bulunamadı.')
    
    doc.add_paragraph('')
    
    # En çok hareket gören ürünler
    doc.add_heading('3. EN ÇOK HAREKET GÖREN ÜRÜNLER', level=1)
    
    top_products = db.session.query(
        Product.marka,
        Product.model,
        db.func.count(Movement.id).label('movement_count')
    ).join(Movement).group_by(
        Product.id, Product.marka, Product.model
    ).order_by(
        db.func.count(Movement.id).desc()
    ).limit(10).all()
    
    if top_products:
        product_table = doc.add_table(rows=1, cols=3)
        product_table.style = 'Table Grid'
        hdr_cells = product_table.rows[0].cells
        hdr_cells[0].text = 'Marka'
        hdr_cells[1].text = 'Model'
        hdr_cells[2].text = 'Hareket Sayısı'
        
        for product in top_products:
            row_cells = product_table.add_row().cells
            row_cells[0].text = product.marka
            row_cells[1].text = product.model
            row_cells[2].text = str(product.movement_count)
    else:
        doc.add_paragraph('Hareket verisi bulunamadı.')
    
    doc.add_paragraph('')
    
    # Marka dağılımı
    doc.add_heading('4. MARKA DAĞILIMI', level=1)
    
    brand_distribution = db.session.query(
        Product.marka,
        db.func.count(Product.id).label('product_count'),
        db.func.sum(Product.adet).label('total_quantity')
    ).group_by(Product.marka).order_by(
        db.func.count(Product.id).desc()
    ).limit(10).all()
    
    if brand_distribution:
        brand_table = doc.add_table(rows=1, cols=3)
        brand_table.style = 'Table Grid'
        hdr_cells = brand_table.rows[0].cells
        hdr_cells[0].text = 'Marka'
        hdr_cells[1].text = 'Ürün Sayısı'
        hdr_cells[2].text = 'Toplam Stok'
        
        for brand in brand_distribution:
            row_cells = brand_table.add_row().cells
            row_cells[0].text = brand.marka
            row_cells[1].text = str(brand.product_count)
            row_cells[2].text = str(brand.total_quantity)
    else:
        doc.add_paragraph('Marka verisi bulunamadı.')
    
    # Dosyayı kaydet
    filename = f"stok_raporu_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join('static', 'reports', filename)
    
    # Reports klasörünü oluştur
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    
    doc.save(filepath)
    return filepath, filename

def create_pdf_report():
    """PDF formatında rapor oluştur"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    story = []
    
    # Türkçe karakter desteği için font kaydet
    try:
        # Windows'ta bulunan Arial fontunu kullan
        pdfmetrics.registerFont(TTFont('Arial', 'C:/Windows/Fonts/arial.ttf'))
        pdfmetrics.registerFont(TTFont('Arial-Bold', 'C:/Windows/Fonts/arialbd.ttf'))
        font_name = 'Arial'
        font_bold = 'Arial-Bold'
    except:
        # Arial bulunamazsa varsayılan font kullan
        font_name = 'Helvetica'
        font_bold = 'Helvetica-Bold'
    
    # Stil tanımlamaları
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontName=font_bold,
        fontSize=18,
        spaceAfter=30,
        alignment=TA_CENTER,
        textColor=colors.darkblue
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontName=font_bold,
        fontSize=14,
        spaceAfter=12,
        textColor=colors.darkblue
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=10
    )
    
    # Başlık
    story.append(Paragraph('ASES STOK TAKİP SİSTEMİ', title_style))
    story.append(Paragraph('ANALİZ RAPORU', title_style))
    story.append(Paragraph(f'Rapor Tarihi: {datetime.now().strftime("%d.%m.%Y %H:%M")}', normal_style))
    story.append(Spacer(1, 20))
    
    # Genel Bilgiler
    story.append(Paragraph('1. GENEL BİLGİLER', heading_style))
    
    # İstatistikler
    total_products = Product.query.count()
    total_quantity = db.session.query(db.func.sum(Product.adet)).scalar() or 0
    avg_stock = total_quantity / total_products if total_products > 0 else 0
    
    high_stock = Product.query.filter(Product.adet > 50).count()
    medium_stock = Product.query.filter(Product.adet.between(10, 50)).count()
    low_stock = Product.query.filter(Product.adet.between(1, 9)).count()
    no_stock = Product.query.filter(Product.adet == 0).count()
    
    stats_data = [
        ['Kategori', 'Değer'],
        ['Toplam Ürün Sayısı', str(total_products)],
        ['Toplam Stok Miktarı', str(total_quantity)],
        ['Ortalama Stok', f"{avg_stock:.2f}"],
        ['Yüksek Stok (>50)', str(high_stock)],
        ['Orta Stok (10-50)', str(medium_stock)],
        ['Düşük Stok (1-9)', str(low_stock)],
        ['Stok Yok (0)', str(no_stock)]
    ]
    
    stats_table = Table(stats_data)
    stats_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), font_bold),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('FONTNAME', (0, 1), (-1, -1), font_name),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(stats_table)
    story.append(Spacer(1, 20))
    
    # Son 7 günlük hareketler
    story.append(Paragraph('2. SON 7 GÜNLÜK HAREKETLER', heading_style))
    
    today = datetime.now().date()
    week_ago = today - timedelta(days=7)
    
    movements = db.session.query(
        db.func.date(Movement.tarih).label('date'),
        Movement.movement_type,
        db.func.sum(Movement.adet).label('total')
    ).filter(
        Movement.tarih >= week_ago
    ).group_by(
        db.func.date(Movement.tarih),
        Movement.movement_type
    ).order_by(db.func.date(Movement.tarih)).all()
    
    if movements:
        movement_data = [['Tarih', 'Hareket Türü', 'Adet']]
        for movement in movements:
            movement_data.append([
                str(movement.date),
                'Giriş' if movement.movement_type == 'giris' else 'Çıkış',
                str(movement.total)
            ])
        
        movement_table = Table(movement_data)
        movement_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), font_bold),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('FONTNAME', (0, 1), (-1, -1), font_name),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(movement_table)
    else:
        story.append(Paragraph('Son 7 günde hareket bulunamadı.', normal_style))
    
    story.append(Spacer(1, 20))
    
    # En çok hareket gören ürünler
    story.append(Paragraph('3. EN ÇOK HAREKET GÖREN ÜRÜNLER', heading_style))
    
    top_products = db.session.query(
        Product.marka,
        Product.model,
        db.func.count(Movement.id).label('movement_count')
    ).join(Movement).group_by(
        Product.id, Product.marka, Product.model
    ).order_by(
        db.func.count(Movement.id).desc()
    ).limit(10).all()
    
    if top_products:
        product_data = [['Marka', 'Model', 'Hareket Sayısı']]
        for product in top_products:
            product_data.append([
                product.marka,
                product.model,
                str(product.movement_count)
            ])
        
        product_table = Table(product_data)
        product_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), font_bold),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('FONTNAME', (0, 1), (-1, -1), font_name),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(product_table)
    else:
        story.append(Paragraph('Hareket verisi bulunamadı.', normal_style))
    
    # PDF'i oluştur
    doc.build(story)
    buffer.seek(0)
    
    # Dosyayı kaydet
    filename = f"stok_raporu_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    filepath = os.path.join('static', 'reports', filename)
    
    # Reports klasörünü oluştur
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    
    with open(filepath, 'wb') as f:
        f.write(buffer.getvalue())
    
    return filepath, filename

def send_report_to_telegram(filepath, filename):
    """Raporu Telegram'a gönder"""
    try:
        telegram_id = int(TELEGRAM_ID)
        token = TELEGRAM_TOKEN
        
        if not telegram_id or not token:
            print("Telegram credentials not found")
            return False
            
        bot = telepot.Bot(token)
        
        # Dosya uzantısına göre mesaj hazırla (Markdown formatı kullanmadan)
        if filename.endswith('.pdf'):
            message_text = f"📊 STOK ANALİZ RAPORU (PDF)\n"
        else:
            message_text = f"📊 STOK ANALİZ RAPORU (WORD)\n"
        
        message_text += f"📅 Tarih: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n"
        message_text += f"📁 Dosya: {filename}\n"
        message_text += f"💻 Cihaz: {socket.gethostname()}\n"
        message_text += f"👤 Kullanıcı: {current_user.username}\n"
        
        # Dosyayı gönder (parse_mode olmadan)
        with open(filepath, 'rb') as f:
            bot.sendDocument(telegram_id, f, caption=message_text)
        
        print(f"Rapor Telegram'a gönderildi: {filename}")
        return True
        
    except Exception as e:
        print(f"Error sending report to Telegram: {e}")
        return False

# Rapor oluşturma ve gönderme route'u
@app.route('/export_report')
@login_required
def export_report():
    """Rapor oluştur ve Telegram'a gönder"""
    try:
        format_type = request.args.get('format', 'pdf')  # pdf veya word
        
        if format_type == 'word':
            filepath, filename = create_word_report()
        else:
            filepath, filename = create_pdf_report()
        
        # Telegram'a gönder
        success = send_report_to_telegram(filepath, filename)
        
        if success:
            flash(f'{filename} raporu başarıyla oluşturuldu ve Telegram\'a gönderildi!', 'success')
        else:
            flash(f'{filename} raporu oluşturuldu ancak Telegram\'a gönderilemedi!', 'warning')
        
        # Geçici dosyayı sil
        try:
            os.remove(filepath)
        except:
            pass
        
        return redirect(url_for('analytics'))
        
    except Exception as e:
        flash(f'Rapor oluşturulurken hata oluştu: {str(e)}', 'error')
        return redirect(url_for('analytics'))


if __name__ == '__main__':
    with app.app_context():
        db.create_all()
        
    # Günlük stok raporu zamanlayıcısını ayarla (her gün saat 20:00)
    schedule.every().day.at("20:00").do(send_daily_stock_report)
    
    # Zamanlayıcıyı ayrı thread'de çalıştır
    scheduler_thread = threading.Thread(target=run_scheduler, daemon=True)
    scheduler_thread.start()
    
    print("Günlük stok raporu zamanlayıcısı başlatıldı (Her gün 20:00)")
    
    app.run(debug=True)
